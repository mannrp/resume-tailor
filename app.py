import streamlit as st
import google.generativeai as genai
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Any
import json
import re
import io
import requests

# Add docx support
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx not installed. DOCX export will not be available.")

# Add PDF processing support
try:
    import pdfplumber
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    st.warning("PDF processing libraries not installed. ATS analysis will be limited.")

# Page configuration
st.set_page_config(page_title="Resume Tailor", layout="wide")

# Main title
st.title("📄✨ Resume Tailor")
st.info("Upload your LaTeX resume and job description to get a tailored version using Google Gemini AI")

# Define resume structure constraints
@dataclass
class SectionConstraints:
    max_lines: int
    max_words_per_line: int
    required: bool = True

@dataclass
class ResumeStructure:
    contact: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=3, max_words_per_line=20, required=True))
    summary: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=3, max_words_per_line=25, required=False))
    experience: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=8, max_words_per_line=20, required=True))  # 4 jobs * 2 bullets
    skills: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=4, max_words_per_line=15, required=True))
    education: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=2, max_words_per_line=20, required=True))
    projects: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=4, max_words_per_line=20, required=False))  # 2 projects * 2 bullets
    certifications: SectionConstraints = field(default_factory=lambda: SectionConstraints(max_lines=2, max_words_per_line=15, required=False))

def escape_latex_characters(text):
    """Escape special LaTeX characters in text"""
    if not text:
        return text
    
    text = re.sub(r'\\(?!(textbf|textit|underline|&|\$|%|#|_|\{|\}|textasciicircum|textasciitilde|hfill))[a-zA-Z]+', '', text)
    
    latex_chars = {
        '&': '\\&',
        '%': '\\%',
        '$': '\\$',
        '#': '\\#',
        '_': '\\_',
        '{': '\\{',
        '}': '\\}',
        '^': '\\textasciicircum{}',
        '~': '\\textasciitilde{}',
    }
    
    for char, escaped in latex_chars.items():
        text = text.replace(char, escaped)
    
    return text

def to_json_serializable(obj):
    """Convert dataclass objects to JSON-serializable dictionaries"""
    if hasattr(obj, '__dict__'):
        return {key: to_json_serializable(value) for key, value in obj.__dict__.items()}
    elif isinstance(obj, list):
        return [to_json_serializable(item) for item in obj]
    elif isinstance(obj, dict):
        return {key: to_json_serializable(value) for key, value in obj.items()}
    else:
        return obj

@dataclass
class ResumeSection:
    title: str
    lines: List[str] = field(default_factory=list)
    # Using a flexible list of dictionaries for subsections to handle varied formats
    subsections: List[Dict[str, Any]] = field(default_factory=list)

@dataclass
class StructuredResume:
    contact: Optional[ResumeSection] = None
    summary: Optional[ResumeSection] = None
    experience: Optional[ResumeSection] = None
    skills: Optional[ResumeSection] = None
    education: Optional[ResumeSection] = None
    projects: Optional[ResumeSection] = None
    certifications: Optional[ResumeSection] = None

# ... (Keep your API key and sidebar configuration code as is) ...

def parse_latex_into_structure(latex_code: str) -> StructuredResume:
    """Parse LaTeX resume into a structured format with separated dates."""
    
    prompt = f"""
Parse this LaTeX resume into a structured JSON format. It is critical to separate headings (like job titles or degrees) from their corresponding dates and locations.

LaTeX Resume:
{latex_code}

Return a JSON object with this exact structure. For sections with multiple entries (like Experience or Education), use the "subsections" list. Each item in the list should be an object with distinct fields for heading, organization, location, date, and bullet points.

IMPORTANT: Look for these LaTeX patterns to extract data correctly:
- \\resumeSubheading{{title}}{{date}}{{organization}}{{location}}
- \\resumeProjectHeading{{title}}{{date}}
- \\textbf{{title | organization | date}} (older format - split these)

{{
    "contact": {{
        "title": "Contact Information",
        "lines": ["Name", "Email | Phone", "Location | LinkedIn"]
    }},
    "summary": {{
        "title": "Professional Summary",
        "lines": ["Summary line 1", "Summary line 2"]
    }},
    "education": {{
        "title": "Education",
        "subsections": [
            {{
                "heading": "Business Administration Diploma – Finance",
                "organization": "Northern Alberta Institute of Technology",
                "location": "Edmonton, AB",
                "date": "Jan. 2023 -- Apr. 2025",
                "bullets": []
            }}
        ]
    }},
    "experience": {{
        "title": "Experience",
        "subsections": [
            {{
                "heading": "Multi Store Supervisor & Sales Associate",
                "organization": "Mobile Planet",
                "location": "Edmonton, AB",
                "date": "May 2023 -- Present",
                "bullets": ["Bullet point 1", "Bullet point 2"]
            }}
        ]
    }},
    "projects": {{
        "title": "Projects",
        "subsections": [
            {{
                "heading": "Real World Credit-Analysis & Lending Simulation",
                "organization": "",
                "location": "",
                "date": "Sep. 2024 -- Dec. 2024",
                "bullets": ["Bullet point 1", "Bullet point 2"]
            }}
        ]
    }},
    "certifications": {{
        "title": "Certifications & Achievements",
        "subsections": [
            {{
                "heading": "Canadian Securities Course (CSC) -- Completed",
                "organization": "",
                "location": "",
                "date": "Apr. 2025",
                "bullets": ["Bullet point 1"]
            }}
        ]
    }},
    "skills": {{
        "title": "Skills",
        "lines": ["Languages: Python, Java", "Frameworks: React, Django"]
    }}
}}

CRITICAL RULES:
- If a section is missing in the LaTeX, omit its key from the JSON.
- For all sections with subsections (Education, Experience, Projects, Certifications), ALWAYS separate the date into the "date" field. DO NOT merge it with the "heading" or "organization".
- When you see patterns like "Job Title | Company | Date", split them into separate fields.
- Extract ONLY the text content, removing all LaTeX formatting commands.
- Return only valid JSON, with no explanatory text before or after.
"""
    
    try:
        model = genai.GenerativeModel(selected_model)
        response = model.generate_content(prompt)
        
        json_text = response.text.strip().replace('```json', '').replace('```', '')
        data = json.loads(json_text)
        
        # Create resume object from parsed data
        resume_data = {}
        for section_name, section_data in data.items():
            resume_data[section_name] = ResumeSection(
                title=section_data.get("title", section_name.capitalize()),
                lines=section_data.get("lines", []),
                subsections=section_data.get("subsections", [])
            )
            
        return StructuredResume(**resume_data)
        
    except Exception as e:
        st.error(f"Error parsing LaTeX into new structure: {str(e)}")
        st.error(f"LLM Response was: {response.text}")
        return None

def structure_to_latex(structured_resume: StructuredResume, original_latex: str = None) -> str:
    """
    Convert structured resume back to LaTeX format using proper resume template commands.
    """
    if original_latex:
        # Try to preserve the original template structure
        return preserve_original_latex_structure(structured_resume, original_latex)

    # Fallback to standard resume template
    latex_lines = [
        "\\documentclass[letterpaper,11pt]{article}",
        "",
        "\\usepackage{latexsym}",
        "\\usepackage[empty]{fullpage}",
        "\\usepackage{titlesec}",
        "\\usepackage{marvosym}",
        "\\usepackage[usenames,dvipsnames]{color}",
        "\\usepackage{verbatim}",
        "\\usepackage{enumitem}",
        "\\usepackage[hidelinks]{hyperref}",
        "\\usepackage{fancyhdr}",
        "\\usepackage[english]{babel}",
        "\\usepackage{tabularx}",
        "\\input{glyphtounicode}",
        "",
        "\\pagestyle{fancy}",
        "\\fancyhf{}",
        "\\fancyfoot{}",
        "\\renewcommand{\\headrulewidth}{0pt}",
        "\\renewcommand{\\footrulewidth}{0pt}",
        "",
        "\\addtolength{\\oddsidemargin}{-0.5in}",
        "\\addtolength{\\evensidemargin}{-0.5in}",
        "\\addtolength{\\textwidth}{1in}",
        "\\addtolength{\\topmargin}{-0.5in}",
        "\\addtolength{\\textheight}{1.0in}",
        "",
        "\\urlstyle{same}",
        "",
        "\\raggedbottom",
        "\\raggedright",
        "\\setlength{\\tabcolsep}{0in}",
        "",
        "\\titleformat{\\section}{",
        "  \\vspace{-3pt}\\scshape\\raggedright\\large",
        "}{}{0em}{}[\\color{black}\\titlerule \\vspace{-4pt}]",
        "",
        "\\pdfgentounicode=1",
        "",
        "\\newcommand{\\resumeItem}[1]{",
        "  \\item\\small{",
        "    {#1 \\vspace{-1pt}}",
        "  }",
        "}",
        "",
        "\\newcommand{\\resumeSubheading}[4]{",
        "  \\vspace{-1pt}\\item",
        "    \\begin{tabular*}{0.97\\textwidth}[t]{l@{\\extracolsep{\\fill}}r}",
        "      \\textbf{#1} & #2 \\\\",
        "      \\textit{\\small#3} & \\textit{\\small #4} \\\\",
        "    \\end{tabular*}\\vspace{-6pt}",
        "}",
        "",
        "\\newcommand{\\resumeProjectHeading}[2]{",
        "    \\item",
        "    \\begin{tabular*}{0.97\\textwidth}{l@{\\extracolsep{\\fill}}r}",
        "      \\small#1 & #2 \\\\",
        "    \\end{tabular*}\\vspace{-6pt}",
        "}",
        "",
        "\\newcommand{\\resumeSubHeadingListStart}{\\begin{itemize}[leftmargin=0.15in, label={}]}",
        "\\newcommand{\\resumeSubHeadingListEnd}{\\end{itemize}}",
        "\\newcommand{\\resumeItemListStart}{\\begin{itemize}}",
        "\\newcommand{\\resumeItemListEnd}{\\end{itemize}\\vspace{-3pt}}",
        "",
        "\\begin{document}",
        ""
    ]

    # Contact section
    if structured_resume.contact and structured_resume.contact.lines:
        latex_lines.append("\\begin{center}")
        for i, line in enumerate(structured_resume.contact.lines):
            escaped_line = escape_latex_characters(line)
            if i == 0:
                latex_lines.append(f"    \\textbf{{\\Huge \\scshape {escaped_line}}} \\\\ \\vspace{{1pt}}")
            else:
                latex_lines.append(f"    \\small {escaped_line}")
        latex_lines.append("\\end{center}")
        latex_lines.append("")

    # Summary section  
    if structured_resume.summary and structured_resume.summary.lines:
        latex_lines.append("\\section{Professional Summary}")
        for line in structured_resume.summary.lines:
            escaped_line = escape_latex_characters(line)
            latex_lines.append(f"\\small{{{escaped_line}}}")
        latex_lines.append("")

    # Education section
    if structured_resume.education and structured_resume.education.subsections:
        latex_lines.append("\\section{Education}")
        latex_lines.append("  \\resumeSubHeadingListStart")
        for item in structured_resume.education.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            organization = escape_latex_characters(item.get("organization", ""))
            location = escape_latex_characters(item.get("location", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            latex_lines.append(f"    \\resumeSubheading")
            latex_lines.append(f"      {{{organization}}}{{{location}}}")
            latex_lines.append(f"      {{{heading}}}{{{date}}}")
            
            bullets = item.get("bullets", [])
            non_empty_bullets = [b for b in bullets if b.strip()]
            if non_empty_bullets:
                latex_lines.append("  \\resumeItemListStart")
                for bullet in non_empty_bullets:
                    latex_lines.append(f"  \\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("  \\resumeItemListEnd")
        latex_lines.append("  \\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Experience section
    if structured_resume.experience and structured_resume.experience.subsections:
        latex_lines.append("\\section{Experience}")
        latex_lines.append("  \\resumeSubHeadingListStart")
        for item in structured_resume.experience.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            organization = escape_latex_characters(item.get("organization", ""))
            location = escape_latex_characters(item.get("location", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            latex_lines.append(f"    \\resumeSubheading")
            latex_lines.append(f"      {{{heading}}}{{{date}}}")
            latex_lines.append(f"      {{{organization}}}{{{location}}}")
            
            bullets = item.get("bullets", [])
            non_empty_bullets = [b for b in bullets if b.strip()]
            if non_empty_bullets:
                latex_lines.append("  \\resumeItemListStart")
                for bullet in non_empty_bullets:
                    latex_lines.append(f"  \\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("  \\resumeItemListEnd")
        latex_lines.append("  \\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Projects section
    if structured_resume.projects and structured_resume.projects.subsections:
        latex_lines.append("\\section{Projects}")
        latex_lines.append("  \\resumeSubHeadingListStart")
        for item in structured_resume.projects.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            latex_lines.append(f"    \\resumeProjectHeading")
            latex_lines.append(f"      {{\\textbf{{{heading}}}}}{{{date}}}")
            
            bullets = item.get("bullets", [])
            non_empty_bullets = [b for b in bullets if b.strip()]
            if non_empty_bullets:
                latex_lines.append("     \\resumeItemListStart")
                for bullet in non_empty_bullets:
                    latex_lines.append(f"  \\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("\\resumeItemListEnd")
        latex_lines.append("  \\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Certifications section
    if structured_resume.certifications and structured_resume.certifications.subsections:
        latex_lines.append("\\section{Certifications \\& Achievements}")
        latex_lines.append("  \\resumeSubHeadingListStart")
        for item in structured_resume.certifications.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            latex_lines.append(f"    \\resumeProjectHeading")
            latex_lines.append(f"      {{\\textbf{{{heading}}}}}{{{date}}}")
            
            bullets = item.get("bullets", [])
            non_empty_bullets = [b for b in bullets if b.strip()]
            if non_empty_bullets:
                latex_lines.append("      \\resumeItemListStart")
                for bullet in non_empty_bullets:
                    latex_lines.append(f"        \\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("      \\resumeItemListEnd")
        latex_lines.append("  \\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Skills section
    if structured_resume.skills and structured_resume.skills.lines:
        latex_lines.append("\\section{Skills}")
        latex_lines.append("  \\begin{itemize}[leftmargin=0.15in, label={}]")
        latex_lines.append("    \\small{\\item{")
        for i, line in enumerate(structured_resume.skills.lines):
            escaped_line = escape_latex_characters(line)
            if i < len(structured_resume.skills.lines) - 1:
                latex_lines.append(f"      {escaped_line} " + r"\\")
            else:
                latex_lines.append(f"      {escaped_line}")
        latex_lines.append("    }}")
        latex_lines.append("  \\end{itemize}")
        latex_lines.append("")
        latex_lines.append("  }}")
        latex_lines.append("\\end{itemize}")
        latex_lines.append("")

    latex_lines.append("\\end{document}")
    return "\n".join(latex_lines)


def structure_to_docx(structured_resume: StructuredResume) -> bytes:
    """Convert structured resume to DOCX format"""
    if not DOCX_AVAILABLE:
        return None
        
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Contact section
    if structured_resume.contact and structured_resume.contact.lines:
        for line in structured_resume.contact.lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Summary section
    if structured_resume.summary and structured_resume.summary.lines:
        doc.add_heading('Professional Summary', level=1)
        for line in structured_resume.summary.lines:
            doc.add_paragraph(line)
        doc.add_paragraph()
    
    # Helper for sections with subsections
    def add_docx_section(title, section_obj):
        if section_obj and section_obj.subsections:
            doc.add_heading(title, level=1)
            for item in section_obj.subsections:
                heading = item.get("heading", "")
                organization = item.get("organization", "")
                location = item.get("location", "")
                date = item.get("date", "")

                # Create a paragraph for the heading line
                p = doc.add_paragraph()
                
                # For Experience and Education: add heading first, then org/location
                if title in ['Experience', 'Education']:
                    p.add_run(heading).bold = True
                    p.add_run("\t")
                    p.add_run(date).italic = True
                    
                    if organization or location:
                        org_line = doc.add_paragraph()
                        org_info = organization
                        if location:
                            org_info += f" | {location}" if organization else location
                        org_line.add_run(org_info).italic = True
                else:
                    # For Projects and Certifications: just heading and date
                    p.add_run(heading).bold = True
                    p.add_run("\t")
                    p.add_run(date).italic = True
                
                # Set up right-aligned tab stop for dates
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                # Add bullet points
                bullets = item.get("bullets", [])
                for bullet in bullets:
                    if bullet.strip():
                        doc.add_paragraph(bullet, style='List Bullet')
            doc.add_paragraph()

    # Add all sections with subsections
    add_docx_section('Education', structured_resume.education)
    add_docx_section('Experience', structured_resume.experience)
    add_docx_section('Projects', structured_resume.projects)
    add_docx_section('Certifications & Achievements', structured_resume.certifications)

    # Skills section
    if structured_resume.skills and structured_resume.skills.lines:
        doc.add_heading('Skills', level=1)
        for line in structured_resume.skills.lines:
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        doc.add_paragraph()

    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()



# API key configuration
api_key = None
try:
    api_key = st.secrets["GEMINI_API_KEY"]
except:
    api_key = st.sidebar.text_input("Enter your Gemini API Key", type="password", help="Get your API key from https://aistudio.google.com/app/apikey")

if api_key:
    genai.configure(api_key=api_key)
else:
    st.error("Please provide your Gemini API key to continue.")

# Configuration options in sidebar
st.sidebar.header("⚙️ Configuration")

# Model selection
selected_model = st.sidebar.selectbox(
    "Choose Gemini Model",
    ["gemini-2.5-flash", "gemini-2.5-flash-lite"],
    index=0,
    help="gemini-1.5-flash is recommended for best results. Use gemini-1.5-pro for higher quality but slower processing."
)

# Structure enforcement options
enforce_structure = st.sidebar.checkbox("📐 Enforce strict structure", value=True, help="Enforces exact line and word limits for each section")
check_ats = st.sidebar.checkbox("🎯 Check ATS compatibility", help="Analyze resume for ATS compatibility after generation")
auto_improve_ats = st.sidebar.checkbox("🔄 Enable ATS optimization (post-review)", help="Adds an option to optimize your resume after reviewing the first version")
generate_docx = st.sidebar.checkbox("📄 Generate DOCX output", help="Create a Word document version", disabled=not DOCX_AVAILABLE)

# Show structure constraints
if enforce_structure:
    with st.sidebar.expander("📋 Structure Constraints"):
        structure = ResumeStructure()
        st.write("**Section Limits:**")
        st.write(f"• Contact: {structure.contact.max_lines} lines, {structure.contact.max_words_per_line} words/line")
        st.write(f"• Experience: {structure.experience.max_lines} lines, {structure.experience.max_words_per_line} words/line")
        st.write(f"• Skills: {structure.skills.max_lines} lines, {structure.skills.max_words_per_line} words/line")
        st.write(f"• Education: {structure.education.max_lines} lines, {structure.education.max_words_per_line} words/line")
        st.write(f"• Projects: {structure.projects.max_lines} lines, {structure.projects.max_words_per_line} words/line")





def apply_ats_feedback(structured_resume: StructuredResume, ats_feedback: str, job_description: str) -> StructuredResume:
    """Apply ATS feedback to improve specific lines without changing structure"""
    
    prompt = f"""
Apply this ATS feedback to improve the resume content line by line, WITHOUT changing the structure.

ATS Feedback:
{ats_feedback}

Job Description:
{job_description}

Current Resume Structure:
{json.dumps(to_json_serializable(structured_resume), indent=2)}

CRITICAL RULES:
1. ONLY edit the content of existing lines - do NOT add or remove lines
2. ONLY edit the text within existing subsections - do NOT add or remove subsections
3. Focus on adding missing keywords mentioned in ATS feedback
4. Improve quantified achievements where suggested
5. Enhance action verbs and impact statements
6. Maintain the same word count limits as the original lines
7. Keep the same heading, organization, location, date structure for subsections

Return the improved structure in the same JSON format with IDENTICAL structure but improved content.
Return only valid JSON, no additional text.
"""
    
    try:
        model = genai.GenerativeModel(selected_model)
        response = model.generate_content(prompt)
        
        # Parse the improved structure
        json_text = response.text.strip()
        if json_text.startswith('```json'):
            json_text = json_text[7:]
        if json_text.endswith('```'):
            json_text = json_text[:-3]
        
        data = json.loads(json_text.strip())
        
        # Convert back to StructuredResume
        improved_resume = StructuredResume()
        
        if "contact" in data:
            improved_resume.contact = ResumeSection(
                title=data["contact"]["title"],
                lines=data["contact"]["lines"]
            )
            
        if "summary" in data:
            improved_resume.summary = ResumeSection(
                title=data["summary"]["title"],
                lines=data["summary"]["lines"]
            )
            
        if "experience" in data:
            improved_resume.experience = ResumeSection(
                title=data["experience"]["title"],
                lines=[],
                subsections=data["experience"].get("subsections", [])
            )
            
        if "skills" in data:
            improved_resume.skills = ResumeSection(
                title=data["skills"]["title"],
                lines=data["skills"]["lines"]
            )
            
        if "education" in data:
            improved_resume.education = ResumeSection(
                title=data["education"]["title"],
                lines=[],
                subsections=data["education"].get("subsections", [])
            )
            
        if "projects" in data:
            improved_resume.projects = ResumeSection(
                title=data["projects"]["title"],
                lines=[],
                subsections=data["projects"].get("subsections", [])
            )
            
        if "certifications" in data:
            improved_resume.certifications = ResumeSection(
                title=data["certifications"]["title"],
                lines=[],
                subsections=data["certifications"].get("subsections", [])
            )
            
        return improved_resume
        
    except Exception as e:
        st.error(f"Error applying ATS feedback: {str(e)}")
        return structured_resume

def optimize_content_lines(structured_resume: StructuredResume, job_description: str, structure_constraints: Optional[ResumeStructure] = None) -> StructuredResume:
    """Optimize resume content for the specific job description while maintaining structure"""
    
    # Convert structured resume to JSON for AI processing
    resume_json = to_json_serializable(structured_resume)
    
    # Build constraint information if provided
    constraint_info = ""
    if structure_constraints and enforce_structure:
        constraint_info = f"""
STRICT STRUCTURE CONSTRAINTS (MUST BE FOLLOWED):
- Contact: Max {structure_constraints.contact.max_lines} lines, {structure_constraints.contact.max_words_per_line} words per line
- Experience: Max {structure_constraints.experience.max_lines} lines total, {structure_constraints.experience.max_words_per_line} words per line
- Skills: Max {structure_constraints.skills.max_lines} lines, {structure_constraints.skills.max_words_per_line} words per line
- Education: Max {structure_constraints.education.max_lines} lines total, {structure_constraints.education.max_words_per_line} words per line
- Projects: Max {structure_constraints.projects.max_lines} lines total, {structure_constraints.projects.max_words_per_line} words per line
- Certifications: Max {structure_constraints.certifications.max_lines} lines total, {structure_constraints.certifications.max_words_per_line} words per line

CRITICAL: Do NOT add or remove sections, subsections, or lines. Only modify the text content within existing lines and bullet points.
"""
    
    prompt = f"""
Tailor this resume for the specific job description. Focus on optimizing content to match job requirements while maintaining the exact same structure.

Job Description:
{job_description}

Current Resume Structure:
{json.dumps(resume_json, indent=2)}

{constraint_info}

INSTRUCTIONS:
1. Analyze the job description for key requirements, skills, and keywords
2. Optimize bullet points and descriptions to highlight relevant experience
3. Use action verbs and quantified achievements where possible
4. Include relevant keywords naturally in the content
5. Maintain the same number of sections, subsections, and lines
6. Keep the same headings, organizations, locations, and dates
7. Only modify the text content of bullet points and description lines

CRITICAL RULES:
- Do NOT add or remove any sections, subsections, or lines
- Do NOT change headings, organization names, locations, or dates
- Do NOT change the structure - only improve the content
- Ensure all bullet points are relevant to the target job
- Use strong action verbs and quantified results
- Include keywords from the job description naturally

Return the improved resume in the same JSON format with identical structure but optimized content.
Return only valid JSON, no additional text.
"""
    
    try:
        model = genai.GenerativeModel(selected_model)
        response = model.generate_content(prompt)
        
        # Parse the response
        json_text = response.text.strip()
        if json_text.startswith('```json'):
            json_text = json_text[7:]
        if json_text.endswith('```'):
            json_text = json_text[:-3]
        
        data = json.loads(json_text.strip())
        
        # Convert back to StructuredResume
        optimized_resume = StructuredResume()
        
        if "contact" in data:
            optimized_resume.contact = ResumeSection(
                title=data["contact"]["title"],
                lines=data["contact"]["lines"]
            )
            
        if "summary" in data:
            optimized_resume.summary = ResumeSection(
                title=data["summary"]["title"],
                lines=data["summary"]["lines"]
            )
            
        if "experience" in data:
            optimized_resume.experience = ResumeSection(
                title=data["experience"]["title"],
                lines=[],
                subsections=data["experience"].get("subsections", [])
            )
            
        if "skills" in data:
            optimized_resume.skills = ResumeSection(
                title=data["skills"]["title"],
                lines=data["skills"]["lines"]
            )
            
        if "education" in data:
            optimized_resume.education = ResumeSection(
                title=data["education"]["title"],
                lines=[],
                subsections=data["education"].get("subsections", [])
            )
            
        if "projects" in data:
            optimized_resume.projects = ResumeSection(
                title=data["projects"]["title"],
                lines=[],
                subsections=data["projects"].get("subsections", [])
            )
            
        if "certifications" in data:
            optimized_resume.certifications = ResumeSection(
                title=data["certifications"]["title"],
                lines=[],
                subsections=data["certifications"].get("subsections", [])
            )
            
        return optimized_resume
        
    except Exception as e:
        st.error(f"Error optimizing resume content: {str(e)}")
        st.error(f"AI Response was: {response.text}")
        return structured_resume

def structure_to_latex(structured_resume: StructuredResume, original_latex: str = None) -> str:
    """
    Convert structured resume back to LaTeX format using proper resume template commands.
    This version corrects the logic to ensure all sections with sub-items
    (Experience, Projects, etc.) use the correct, separated formatting.
    """
    # If original_latex is provided, use a replacement strategy.
    if original_latex:
        return preserve_original_latex_structure(structured_resume, original_latex)

    # --- Fallback Generator using a Standard Template ---
    # This part contains the corrected logic for generating from scratch.
    latex_lines = [
        "\\documentclass[letterpaper,11pt]{article}",
        "\\usepackage{latexsym}",
        "\\usepackage[empty]{fullpage}",
        "\\usepackage{titlesec}",
        "\\usepackage{marvosym}",
        "\\usepackage[usenames,dvipsnames]{color}",
        "\\usepackage{verbatim}",
        "\\usepackage{enumitem}",
        "\\usepackage[hidelinks]{hyperref}",
        "\\usepackage{fancyhdr}",
        "\\usepackage[english]{babel}",
        "\\usepackage{tabularx}",
        "\\input{glyphtounicode}",
        "",
        "\\pagestyle{fancy}",
        "\\fancyhf{}",
        "\\fancyfoot{}",
        "\\renewcommand{\\headrulewidth}{0pt}",
        "\\renewcommand{\\footrulewidth}{0pt}",
        "",
        "\\addtolength{\\oddsidemargin}{-0.5in}",
        "\\addtolength{\\evensidemargin}{-0.5in}",
        "\\addtolength{\\textwidth}{1in}",
        "\\addtolength{\\topmargin}{-0.5in}",
        "\\addtolength{\\textheight}{1.0in}",
        "",
        "\\urlstyle{same}",
        "",
        "\\raggedbottom",
        "\\raggedright",
        "\\setlength{\\tabcolsep}{0in}",
        "",
        "\\titleformat{\\section}{",
        "  \\vspace{-3pt}\\scshape\\raggedright\\large",
        "}{}{0em}{}[\\color{black}\\titlerule \\vspace{-4pt}]",
        "",
        "\\pdfgentounicode=1",
        "",
        "\\newcommand{\\resumeItem}[1]{",
        "  \\item\\small{",
        "    {#1 \\vspace{-1pt}}",
        "  }",
        "}",
        "",
        "\\newcommand{\\resumeSubheading}[4]{",
        "  \\vspace{-1pt}\\item",
        "    \\begin{tabular*}{0.97\\textwidth}[t]{l@{\\extracolsep{\\fill}}r}",
        "      \\textbf{#1} & #2 \\\\",
        "      \\textit{\\small#3} & \\textit{\\small #4} \\\\",
        "    \\end{tabular*}\\vspace{-6pt}",
        "}",
        "",
        "\\newcommand{\\resumeProjectHeading}[2]{",
        "    \\item",
        "    \\begin{tabular*}{0.97\\textwidth}{l@{\\extracolsep{\\fill}}r}",
        "      \\small#1 & #2 \\\\",
        "    \\end{tabular*}\\vspace{-6pt}",
        "}",
        "",
        "\\newcommand{\\resumeSubHeadingListStart}{\\begin{itemize}[leftmargin=0.15in, label={}]}",
        "\\newcommand{\\resumeSubHeadingListEnd}{\\end{itemize}}",
        "\\newcommand{\\resumeItemListStart}{\\begin{itemize}}",
        "\\newcommand{\\resumeItemListEnd}{\\end{itemize}\\vspace{-3pt}}",
        "",
        "\\begin{document}",
        ""
    ]

    # Contact section (remains the same)
    if structured_resume.contact and structured_resume.contact.lines:
        latex_lines.append("\\begin{center}")
        for i, line in enumerate(structured_resume.contact.lines):
            escaped_line = escape_latex_characters(line)
            if i == 0:
                latex_lines.append(f"    \\textbf{{\\Huge \\scshape {escaped_line}}} \\\\ \\vspace{{1pt}}")
            else:
                latex_lines.append(f"    \\small {escaped_line}")
        latex_lines.append("\\end{center}")
        latex_lines.append("")

    # Summary section (remains the same)
    if structured_resume.summary and structured_resume.summary.lines:
        latex_lines.append("\\section{Professional Summary}")
        for line in structured_resume.summary.lines:
            latex_lines.append(f"\\small{{{escape_latex_characters(line)}}}")
        latex_lines.append("")

    # --- CORRECTED SECTIONS ---

    # Education section
    if structured_resume.education and structured_resume.education.subsections:
        latex_lines.append("\\section{Education}")
        latex_lines.append("  \\\\resumeSubHeadingListStart")
        for item in structured_resume.education.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            organization = escape_latex_characters(item.get("organization", ""))
            location = escape_latex_characters(item.get("location", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            # Corrected argument order: Heading, Date, Organization, Location
            latex_lines.append(f"    \\\\resumeSubheading{{{heading}}}{{{date}}}{{{organization}}}{{{location}}}")
            
            # Bullet points (if any)
            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                latex_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    latex_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("    \\\\resumeItemListEnd")
        latex_lines.append("  \\\\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Experience section
    if structured_resume.experience and structured_resume.experience.subsections:
        latex_lines.append("\\section{Experience}")
        latex_lines.append("  \\\\resumeSubHeadingListStart")
        for item in structured_resume.experience.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            organization = escape_latex_characters(item.get("organization", ""))
            location = escape_latex_characters(item.get("location", ""))
            date = escape_latex_characters(item.get("date", ""))

            # Uses the same \resumeSubheading command for consistency
            latex_lines.append(f"    \\\\resumeSubheading{{{heading}}}{{{date}}}{{{organization}}}{{{location}}}")
            
            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                latex_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    latex_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("    \\\\resumeItemListEnd")
        latex_lines.append("  \\\\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Projects section
    if structured_resume.projects and structured_resume.projects.subsections:
        latex_lines.append("\\section{Projects}")
        latex_lines.append("  \\\\resumeSubHeadingListStart")
        for item in structured_resume.projects.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            date = escape_latex_characters(item.get("date", ""))
            
            # Uses the \resumeProjectHeading command
            latex_lines.append(f"    \\\\resumeProjectHeading{{\\\\textbf{{{heading}}}}}{{{date}}}")

            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                latex_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    latex_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("    \\\\resumeItemListEnd")
        latex_lines.append("  \\\\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Certifications section
    if structured_resume.certifications and structured_resume.certifications.subsections:
        latex_lines.append("\\section{Certifications \\& Achievements}")
        latex_lines.append("  \\\\resumeSubHeadingListStart")
        for item in structured_resume.certifications.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            date = escape_latex_characters(item.get("date", ""))

            # Uses the \resumeProjectHeading command
            latex_lines.append(f"    \\\\resumeProjectHeading{{\\\\textbf{{{heading}}}}}{{{date}}}")

            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                latex_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    latex_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                latex_lines.append("    \\\\resumeItemListEnd")
        latex_lines.append("  \\\\resumeSubHeadingListEnd")
        latex_lines.append("")

    # Skills section (remains the same)
    if structured_resume.skills and structured_resume.skills.lines:
        latex_lines.append("\\section{Skills}")
        latex_lines.append("  \\begin{itemize}[leftmargin=0.15in, label={}]")
        latex_lines.append("    \\small{\\item{")
        for i, line in enumerate(structured_resume.skills.lines):
            escaped_line = escape_latex_characters(line)
            if i < len(structured_resume.skills.lines) - 1:
                latex_lines.append(f"      {escaped_line} " + r"\\")
            else:
                latex_lines.append(f"      {escaped_line}")
        latex_lines.append("    }}")
        latex_lines.append("  \\end{itemize}")
        latex_lines.append("")

    latex_lines.append("\\end{document}")
    return "\n".join(latex_lines)


def preserve_original_latex_structure(structured_resume: StructuredResume, original_latex: str) -> str:
    """
    Preserve the original LaTeX template structure while updating content.
    This version ensures the correct commands and argument order are used.
    """
    import re
    updated_latex = original_latex

    def replace_section_content(full_latex, section_title, new_section_body):
        # Regex to find a section and replace its content
        pattern = re.compile(
            r"(\\section\{" + re.escape(section_title) + r"\})(.+?)(?=\\section|\\end\{document\})",
            re.DOTALL | re.IGNORECASE
        )
        # If the section is found, replace its body with the new generated content
        if pattern.search(full_latex):
            return pattern.sub(r"\1" + "\n" + new_section_body, full_latex, 1)
        # If not found, append the new section at the end before \end{document}
        else:
            new_section_text = f"\\section{{{section_title}}}\n{new_section_body}\n\n\\end{{document}}"
            return full_latex.replace("\\end{document}", new_section_text)

    # Experience Section
    if structured_resume.experience and structured_resume.experience.subsections:
        content_lines = ["  \\\\resumeSubHeadingListStart"]
        for item in structured_resume.experience.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            organization = escape_latex_characters(item.get("organization", ""))
            location = escape_latex_characters(item.get("location", ""))
            date = escape_latex_characters(item.get("date", ""))
            content_lines.append(f"    \\\\resumeSubheading{{{heading}}}{{{date}}}{{{organization}}}{{{location}}}")
            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                content_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    content_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                content_lines.append("    \\\\resumeItemListEnd")
        content_lines.append("  \\\\resumeSubHeadingListEnd")
        updated_latex = replace_section_content(updated_latex, "Experience", "\n".join(content_lines))
    
    # Projects Section
    if structured_resume.projects and structured_resume.projects.subsections:
        content_lines = ["  \\\\resumeSubHeadingListStart"]
        for item in structured_resume.projects.subsections:
            heading = escape_latex_characters(item.get("heading", ""))
            date = escape_latex_characters(item.get("date", ""))
            content_lines.append(f"    \\\\resumeProjectHeading{{\\\\textbf{{{heading}}}}}{{{date}}}")
            bullets = [b for b in item.get("bullets", []) if b.strip()]
            if bullets:
                content_lines.append("    \\\\resumeItemListStart")
                for bullet in bullets:
                    content_lines.append(f"      \\\\resumeItem{{{escape_latex_characters(bullet)}}}")
                content_lines.append("    \\\\resumeItemListEnd")
        content_lines.append("  \\\\resumeSubHeadingListEnd")
        updated_latex = replace_section_content(updated_latex, "Projects", "\n".join(content_lines))

    # Add other sections like Education, Certifications, etc., following the same pattern as above...

    return updated_latex


def structure_to_docx(structured_resume: StructuredResume) -> bytes:
    """Convert structured resume to DOCX format"""
    if not DOCX_AVAILABLE:
        return None
        
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Contact section
    if structured_resume.contact and structured_resume.contact.lines:
        for line in structured_resume.contact.lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # Summary section
    if structured_resume.summary and structured_resume.summary.lines:
        doc.add_heading('Professional Summary', level=1)
        for line in structured_resume.summary.lines:
            doc.add_paragraph(line)
        doc.add_paragraph()
    
    # Helper for sections with subsections
    def add_docx_section(title, section_obj):
        if section_obj and section_obj.subsections:
            doc.add_heading(title, level=1)
            for item in section_obj.subsections:
                heading = item.get("heading", "")
                organization = item.get("organization", "")
                location = item.get("location", "")
                date = item.get("date", "")

                # Create a paragraph for the heading line
                p = doc.add_paragraph()
                
                # For Experience and Education: add heading first, then org/location
                if title in ['Experience', 'Education']:
                    p.add_run(heading).bold = True
                    p.add_run("\t")
                    p.add_run(date).italic = True
                    
                    if organization or location:
                        org_line = doc.add_paragraph()
                        org_info = organization
                        if location:
                            org_info += f" | {location}" if organization else location
                        org_line.add_run(org_info).italic = True
                else:
                    # For Projects and Certifications: just heading and date
                    p.add_run(heading).bold = True
                    p.add_run("\t")
                    p.add_run(date).italic = True
                
                # Set up right-aligned tab stop for dates
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                # Add bullet points
                bullets = item.get("bullets", [])
                for bullet in bullets:
                    if bullet.strip():
                        doc.add_paragraph(bullet, style='List Bullet')
            doc.add_paragraph()

    # Add all sections with subsections
    add_docx_section('Education', structured_resume.education)
    add_docx_section('Experience', structured_resume.experience)
    add_docx_section('Projects', structured_resume.projects)
    add_docx_section('Certifications & Achievements', structured_resume.certifications)

    # Skills section
    if structured_resume.skills and structured_resume.skills.lines:
        doc.add_heading('Skills', level=1)
        for line in structured_resume.skills.lines:
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        doc.add_paragraph()

    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

# Function to clean LaTeX output
def clean_latex_output(latex_code):
    """Remove markdown code blocks from LaTeX output"""
    # Remove ```latex at the beginning
    latex_code = re.sub(r'^```latex\s*', '', latex_code, flags=re.MULTILINE)
    # Remove ``` at the end
    latex_code = re.sub(r'\s*```\s*$', '', latex_code, flags=re.MULTILINE)
    return latex_code.strip()

# Function to extract text from PDF for ATS analysis
def extract_text_from_pdf(pdf_bytes):
    """Extract clean text from PDF bytes for ATS analysis"""
    try:
        # Try pdfplumber first (better text extraction)
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception as e:
        st.warning(f"Could not parse PDF with pdfplumber ({e}), falling back to PyPDF2.")
        try:
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e2:
            st.error(f"Failed to extract text from PDF with all methods: {e2}")
            return None

# Function to extract clean text from LaTeX (improved)
def extract_text_from_latex(latex_code):
    """Extract clean text from LaTeX code for ATS analysis"""
    # Remove comments
    text = re.sub(r'%.*', '', latex_code)
    
    # Remove document structure commands
    text = re.sub(r'\\documentclass\{[^}]*\}', '', text)
    text = re.sub(r'\\usepackage\{[^}]*\}', '', text)
    text = re.sub(r'\\begin\{document\}', '', text)
    text = re.sub(r'\\end\{document\}', '', text)
    
    # Remove section commands but keep the text
    text = re.sub(r'\\section\*?\{([^}]*)\}', r'\n\n\1\n', text)
    text = re.sub(r'\\subsection\*?\{([^}]*)\}', r'\n\1\n', text)
    text = re.sub(r'\\subsubsection\*?\{([^}]*)\}', r'\n\1\n', text)
    
    # Remove resume-specific commands but keep content
    text = re.sub(r'\\resumeSubheading\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}', r'\1 | \2 | \3 | \4', text)
    text = re.sub(r'\\resumeSubSubheading\{([^}]*)\}\{([^}]*)\}', r'\1 | \2', text)
    text = re.sub(r'\\resumeProjectHeading\{([^}]*)\}\{([^}]*)\}', r'\1 | \2', text)
    text = re.sub(r'\\resumeSubItem\{([^}]*)\}', r'\1', text)
    
    # Remove other common LaTeX commands but preserve content
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\href\{[^}]*\}\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)
    
    # Remove list environments but keep items
    text = re.sub(r'\\begin\{itemize\}', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\begin\{enumerate\}', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\item\s*', '\n• ', text)
    
    # Remove remaining LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\*?\{[^}]*\}', '', text)
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
    
    # Remove braces and clean up
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n', text)
    
    return text.strip()
# Function to check ATS compatibility
def check_ats_compatibility(resume_text, job_description, is_from_pdf=False):
    """Analyze resume for ATS (Applicant Tracking System) compatibility using AI"""
    try:
        # Use selected model for ATS analysis
        model = genai.GenerativeModel(selected_model)
        
        text_source = "PDF-extracted text" if is_from_pdf else "LaTeX-extracted text"
        
        prompt = f"""
Analyze this resume for ATS (Applicant Tracking System) compatibility and provide a score out of 100 with specific feedback.

NOTE: This text was extracted from {text_source}.

IMPORTANT: When analyzing contact information, look for:
- Name (usually appears first or prominently)
- Email address (contains @ symbol)
- Phone number (various formats)
- Location/address
- LinkedIn profile or website URLs

Resume Text to Analyze:
{resume_text}

Job Description for Keyword Matching:
{job_description}

Provide response in this format:
ATS SCORE: [0-100]

CONTACT INFORMATION ANALYSIS:
- Found contact elements: [list what contact info was detected]
- Missing contact elements: [list what's missing]

STRENGTHS:
- [list 3-5 strengths with specific examples from the resume]

IMPROVEMENTS:
- [list 3-5 specific improvements needed]

KEYWORD ANALYSIS:
- Missing keywords: [list important keywords from job description not found in resume]
- Well-matched keywords: [list keywords that appear in both]
- Keyword match rate: [percentage]

FORMATTING ANALYSIS:
- Section headers detected: [list section headers found]
- Formatting issues: [any problematic formatting]
"""
        
        response = model.generate_content(prompt)
        return response.text
        
    except Exception as e:
        return f"Error analyzing ATS compatibility: {str(e)}"

# Function to compile LaTeX to PDF using online service only
def compile_latex_to_pdf(latex_code):
    """Compile LaTeX code to PDF using online service - more reliable for web apps"""
    
    # Debug: show the LaTeX code to help diagnose errors
    with st.expander("🔍 Debug: Generated LaTeX Code"):
        st.code(latex_code, language="latex")
        
    st.info("🌐 Compiling PDF using online LaTeX service...")
    return compile_latex_to_pdf_online(latex_code)

def compile_latex_to_pdf_online(latex_code):
    """Compile LaTeX code to PDF using an online LaTeX compiler"""
    try:
        url = "https://latex.ytotech.com/builds/sync"
        files = {'file': ('main.tex', latex_code, 'text/plain')}
        headers = {'User-Agent': 'Resume-Tailor-App/1.0'}
        
        response = requests.post(url, files=files, headers=headers, timeout=45)
        
        # Check for a successful response that is a PDF
        if response.status_code in [200, 201] and response.content and response.content.startswith(b'%PDF'):
            st.success("✅ PDF compiled successfully using online LaTeX service!")
            return response.content
        else:
            # --- THIS IS THE CORRECTED ERROR HANDLING ---
            error_msg = f"Online PDF compilation failed. HTTP Status: {response.status_code}"
            try:
                # The service returns error details in a JSON object
                error_details = response.json()
                logs = error_details.get('logs', 'No logs were returned.')
                
                # The log is just a string, NOT base64. We just display it.
                # We can also clean up the "b'...'" wrapper if it exists.
                if isinstance(logs, str) and logs.startswith("b'") and logs.endswith("'"):
                    logs = logs[2:-1]

                error_msg += f"\n\n**Compiler Logs:**\n{logs}"

            except (json.JSONDecodeError, AttributeError):
                # If the response isn't JSON, show the raw text
                error_msg += f"\n\n**Raw Error Response:**\n{response.text}"
            
            st.error(error_msg)
            return None
            
    except requests.exceptions.Timeout:
        st.error("Online compilation timed out. The service might be busy. Please try again.")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred during online compilation: {str(e)}")
        return None

# --- Main App Logic ---
if api_key:
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("📝 Input")
        latex_input = st.text_area(
            "LaTeX Resume Code",
            height=300,
            placeholder="Paste your LaTeX resume code here..."
        )
        
        job_description = st.text_area(
            "Job Description",
            height=150,
            placeholder="Paste the job description here..."
        )
    
    with col2:
        st.header("⚡ Output")
        
        if st.button("🎯 Process Resume", type="primary"):
            if latex_input and job_description:
                st.session_state.clear() # Clear previous runs
                
                with st.spinner("🔍 Parsing resume structure..."):
                    structured_resume = parse_latex_into_structure(latex_input)
                
                if structured_resume:
                    st.success("✅ Resume structure parsed!")
                    
                    with st.expander("📋 Parsed Structure (for review)"):
                        st.json(to_json_serializable(structured_resume))
                    
                    with st.spinner("✨ Tailoring content for the job..."):
                        structure_constraints = ResumeStructure() if enforce_structure else None
                        optimized_resume = optimize_content_lines(structured_resume, job_description, structure_constraints)
                    
                    with st.spinner("📄 Generating tailored LaTeX code..."):
                        final_latex = structure_to_latex(optimized_resume, latex_input)
                        cleaned_latex = clean_latex_output(final_latex)
                    
                    st.session_state['generated_latex'] = cleaned_latex
                    st.session_state['structured_resume'] = optimized_resume
                    st.session_state['original_latex'] = latex_input
                    st.session_state['job_description'] = job_description
                    
                    st.subheader("📋 Generated LaTeX")
                    st.code(cleaned_latex, language='latex')
                    
                    st.download_button(
                        label="💾 Download LaTeX",
                        data=cleaned_latex,
                        file_name="tailored_resume.tex",
                        mime="text/plain"
                    )
                    
                    if check_ats:
                        with st.spinner("🎯 Analyzing for ATS compatibility..."):
                            pdf_bytes = compile_latex_to_pdf(final_latex)
                            if pdf_bytes:
                                resume_text = extract_text_from_pdf(pdf_bytes)
                                if resume_text:
                                    ats_analysis = check_ats_compatibility(resume_text, job_description, is_from_pdf=True)
                                else:
                                    ats_analysis = "Could not extract text from PDF to perform ATS analysis."
                            else:
                                ats_analysis = "PDF compilation failed, so ATS analysis could not be performed on the final output."
                        
                        st.subheader("🎯 ATS Compatibility Analysis")
                        st.text_area("Analysis Result", ats_analysis, height=250)
                        st.session_state['ats_analysis'] = ats_analysis
                else:
                    st.error("❌ Failed to parse resume structure. Please check your LaTeX format.")
            else:
                st.warning("⚠️ Please provide both LaTeX resume code and job description.")

# Post-processing and download section
if 'generated_latex' in st.session_state:
    st.header("📄 Download & Refine")
    
    # Choose which version to compile
    if 'improved_latex' in st.session_state:
        version_choice = st.radio(
            "Choose version to download:",
            ["Initial Tailored Version", "ATS-Improved Version"],
            horizontal=True
        )
        latex_to_compile = st.session_state['improved_latex'] if version_choice == "ATS-Improved Version" else st.session_state['generated_latex']
        filename_prefix = "ats_improved_resume" if version_choice == "ATS-Improved Version" else "tailored_resume"
    else:
        latex_to_compile = st.session_state['generated_latex']
        filename_prefix = "tailored_resume"

    # ATS Optimization Button
    if 'ats_analysis' in st.session_state and auto_improve_ats:
        st.subheader("🔄 Optional: Apply ATS Feedback")
        if st.button("💪 Generate ATS-Improved Version"):
            with st.spinner("Applying ATS feedback..."):
                improved_resume = apply_ats_feedback(st.session_state['structured_resume'], st.session_state['ats_analysis'], st.session_state['job_description'])
                improved_latex = structure_to_latex(improved_resume, st.session_state['original_latex'])
                st.session_state['improved_latex'] = clean_latex_output(improved_latex)
                st.session_state['improved_structured_resume'] = improved_resume
                st.success("✅ ATS-Improved version generated! Select it above to download.")
                st.rerun()

    # Download buttons
    pdf_col, docx_col = st.columns(2)

    with pdf_col:
        if st.button("📄 Compile & Download PDF"):
            with st.spinner("Compiling PDF... This may take a moment."):
                pdf_bytes = compile_latex_to_pdf(latex_to_compile)
                if pdf_bytes:
                    st.download_button(
                        label="✅ Download PDF Now",
                        data=pdf_bytes,
                        file_name=f"{filename_prefix}.pdf",
                        mime="application/pdf",
                        type="primary"
                    )
                else:
                    st.error("PDF compilation failed. Check the debug logs above.")
    
    with docx_col:
        if generate_docx and DOCX_AVAILABLE:
            if st.button("📄 Generate & Download DOCX"):
                structured_data = st.session_state.get('improved_structured_resume') if 'improved_structured_resume' in st.session_state and filename_prefix.startswith("ats") else st.session_state['structured_resume']
                docx_bytes = structure_to_docx(structured_data)
                if docx_bytes:
                    st.download_button(
                        label="✅ Download DOCX Now",
                        data=docx_bytes,
                        file_name=f"{filename_prefix}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

# Help Section
with st.expander("📚 Help & Information"):
    st.write("...") # Keeping help section as is